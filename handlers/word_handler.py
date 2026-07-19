# handlers/word_handler.py
"""
Word文档导出模块 - 使用 python-docx
"""
from pathlib import Path
from typing import List, Dict
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class WordHandler:
    """Word文档处理器"""
    
    def __init__(self):
        self.doc = None
    
    def export_products(self, products: List[Dict], output_path: Path, 
                        category_name: str = "全产品") -> bool:
        try:
            self.doc = Document()
            
            self._add_title(category_name)
            self._add_stats(products)
            self._add_summary_table(products)
            self._add_product_details(products)
            
            self.doc.save(str(output_path))
            print(f"✅ Word文档已生成: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Word导出失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def export_single_product(self, product: Dict, output_path: Path) -> bool:
        """
        导出单个产品到独立的Word文档
        """
        try:
            self.doc = Document()
            
            self._add_single_product(product)
            
            self.doc.save(str(output_path))
            print(f"✅ 单产品Word已生成: {output_path.name}")
            return True
            
        except Exception as e:
            print(f"❌ 单产品Word导出失败: {e}")
            return False
            
    def _add_title(self, category_name: str):
        title = self.doc.add_heading(f"安利日本产品目录 - {category_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        time_para = self.doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
    
    def _add_stats(self, products: List[Dict]):
        total = len(products)
        with_sharebar = sum(1 for p in products if p.get('has_sharebar', False))
        without_sharebar = total - with_sharebar
        
        stats = self.doc.add_paragraph()
        stats.add_run("统计信息\n").bold = True
        stats.add_run(f"   * 产品总数: {total}\n")
        stats.add_run(f"   * 有Sharebar: {with_sharebar}\n")
        stats.add_run(f"   * 无Sharebar: {without_sharebar}")
        self.doc.add_paragraph()
    
    def _add_summary_table(self, products: List[Dict]):
        """汇总表：显示三语名称"""
        self.doc.add_heading("产品汇总表", level=1)
        
        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ['序号', '产品ID', '日文名称', '中文名称', '英文名称', '状态']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, product in enumerate(products, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = product.get('product_id', '') or ''
            row.cells[2].text = (product.get('name_ja') or product.get('name') or '')[:30]
            row.cells[3].text = (product.get('name_zh') or '')[:30]
            row.cells[4].text = (product.get('name_en') or '')[:30]
            row.cells[5].text = '有' if product.get('has_sharebar') else '无'
        
        self.doc.add_paragraph()
    
    def _add_product_details(self, products: List[Dict]):
        """详细信息：显示三语名称 + 大图"""
        self.doc.add_page_break()
        self.doc.add_heading("产品详细信息", level=1)
        
        for idx, product in enumerate(products, 1):
            if idx > 1 and idx % 8 == 1:
                self.doc.add_page_break()
            
            self._add_single_product(product, idx)
    
    def _add_single_product(self, product: Dict, idx: int = None):
        """
        添加单个产品的详细信息（可被独立导出复用）
        """
        # 三语名称
        name_ja = (product.get('name_ja') or product.get('name') or '未知产品')
        name_zh = product.get('name_zh') or ''
        name_en = product.get('name_en') or ''
        product_id = product.get('product_id', '') or ''
        
        # 标题：显示序号 + 日文
        if idx:
            heading = f"{idx}. {name_ja}"
        else:
            heading = f"产品: {name_ja}"
        self.doc.add_heading(heading, level=2)
        
        # 产品ID
        id_para = self.doc.add_paragraph()
        id_para.add_run(f"产品ID: ").bold = True
        id_para.add_run(f"{product_id}")
        
        # 多语言名称区块
        p = self.doc.add_paragraph()
        p.add_run("📝 多语言名称\n").bold = True
        p.add_run(f"   🇯🇵 日文: {name_ja}\n")
        if name_zh:
            p.add_run(f"   🇨🇳 中文: {name_zh}\n")
        if name_en:
            p.add_run(f"   🇬🇧 英文: {name_en}\n")
        
        # 基本信息
        info = self.doc.add_paragraph()
        info.add_run("🔗 基本信息\n").bold = True
        
        sharebar = product.get('sharebar', '') or ''
        info.add_run(f"   Sharebar: {sharebar if sharebar else '无'}\n")
        
        status = '有Sharebar ✅' if product.get('has_sharebar') else '无Sharebar ❌'
        info.add_run(f"   状态: {status}\n")
        
        url = product.get('url', '') or ''
        info.add_run(f"   URL: {url}")
        
        # 图片
        merged_path = product.get('merged_path')
        if merged_path and Path(merged_path).exists():
            try:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(merged_path), width=Inches(5.0))
            except Exception as e:
                self.doc.add_paragraph(f"[图片加载失败: {e}]")
        elif product.get('image_path') and Path(product['image_path']).exists():
            try:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(product['image_path']), width=Inches(4.0))
            except:
                self.doc.add_paragraph("[无图片]")
        else:
            self.doc.add_paragraph("[无图片]")
        
        self.doc.add_paragraph('_' * 60)
        